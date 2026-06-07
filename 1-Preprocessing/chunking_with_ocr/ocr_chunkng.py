# -*- coding: utf-8 -*-
# Streamlit + OCR (Qwen2.5-VL) ile gelişmiş doküman chunking

import streamlit as st
import pandas as pd
from io import BytesIO
import pypdf
import docx
import openpyxl
import chardet
from typing import List, Dict, Any, Optional
import json
import math

# OCR / Görsel işleme
from PIL import Image
import pypdfium2 as pdfium
import torch
from transformers import Qwen2VLForConditionalGeneration, AutoProcessor

# =========================
# Sayfa yapılandırması
# =========================
st.set_page_config(
    page_title="Doküman Chunking (OCR'lı)",
    page_icon="",
    layout="wide"
)

# =========================
# Stil
# =========================
st.markdown("""
    <style>
    .chunk-container {
        background-color: #f0f2f6;
        padding: 15px;
        border-radius: 10px;
        margin-bottom: 10px;
        border-left: 4px solid #4CAF50;
    }
    .chunk-header {
        font-weight: bold;
        color: #2c3e50;
        margin-bottom: 10px;
    }
    .chunk-text {
        font-family: 'Courier New', monospace;
        white-space: pre-wrap;
        word-wrap: break-word;
    }
    .stats-container {
        background-color: #e8f4f8;
        padding: 10px;
        border-radius: 5px;
        margin-bottom: 20px;
    }
    .file-info {
        background-color: #f8f9fa;
        padding: 8px;
        border-radius: 5px;
        margin: 5px 0;
        border-left: 3px solid #007bff;
    }
    </style>
""", unsafe_allow_html=True)

# =========================
# Başlık
# =========================
st.title(" Gelişmiş Doküman Chunking (Qwen2.5-VL OCR)")
st.markdown("### PDF/DOCX/görsel dosyalarınızı OCR ile okuyup anlamlı parçalara bölün")

# =========================
# Session state
# =========================
if 'all_chunks' not in st.session_state:
    st.session_state.all_chunks = []
if 'processed_files' not in st.session_state:
    st.session_state.processed_files = []

# =========================
# Yardımcılar
# =========================
def detect_encoding(file_bytes: bytes) -> str:
    res = chardet.detect(file_bytes)
    return res['encoding'] if res['encoding'] else 'utf-8'

@st.cache_resource(show_spinner=True)
def load_qwen_model():
    """
    Qwen/Qwen2.5-VL-7B-Instruct model ve processor'ü cache ederek yükler.
    GPU varsa device_map='auto' ile GPU'ya atar, yoksa CPU.
    """
    model_id = "Qwen/Qwen2.5-VL-7B-Instruct"
    try:
        model = Qwen2VLForConditionalGeneration.from_pretrained(
            model_id,
            device_map="auto",
            torch_dtype="auto",
            trust_remote_code=True
        )
        processor = AutoProcessor.from_pretrained(model_id, trust_remote_code=True)
        return model, processor
    except Exception as e:
        st.error(f"Qwen modeli yüklenemedi: {e}")
        raise

def qwen_ocr_image(pil_image: Image.Image,
                   prompt: Optional[str] = None,
                   max_new_tokens: int = 1024) -> str:
    """
    Bir görselden OCR metni üretir (Qwen2.5-VL ile).
    """
    model, processor = load_qwen_model()

    if prompt is None or not prompt.strip():
        prompt = (
            "Aşağıdaki görseldeki TÜM metni, doğal okuma sırası ve satır sonlarını "
            "olabildiğince koruyarak çıkar. Sadece ham metni döndür; ek açıklama, "
            "etiket veya format ekleme."
        )

    messages = [
        {
            "role": "user",
            "content": [
                {"type": "image", "image": pil_image},
                {"type": "text",  "text": prompt}
            ]
        }
        # assistant cevabını model üretecek
    ]

    # Qwen chat template
    text = processor.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
    inputs = processor(text=[text], images=[pil_image], return_tensors="pt")
    # modele uygun cihaza taşı
    device = model.device
    inputs = {k: v.to(device) if hasattr(v, "to") else v for k, v in inputs.items()}

    with torch.inference_mode():
        generated_ids = model.generate(
            **inputs,
            max_new_tokens=max_new_tokens,
            do_sample=False,
            eos_token_id=processor.tokenizer.eos_token_id
        )

    out = processor.batch_decode(generated_ids, skip_special_tokens=True)[0]
    # Temizlik - şablon kalıntıları vs.
    for token in ["<|system|>", "<|user|>", "<|assistant|>", "<|im_start|>", "<|im_end|>"]:
        out = out.replace(token, "")
    # Bazı durumlarda 'assistant' kelimesinden sonra metin gelebilir
    low = out.lower()
    if "assistant" in low:
        idx = low.rfind("assistant")
        out = out[idx + len("assistant"):].lstrip(": \n")

    return out.strip()

def render_pdf_page(pdf_doc: pdfium.PdfDocument, page_index: int, scale: float = 2.0) -> Image.Image:
    page = pdf_doc[page_index]
    pil = page.render(scale=scale).to_pil()
    return pil.convert("RGB")

# =========================
# Reader'lar (OCR destekli)
# =========================
def read_pdf_ocr(file, ocr_enabled: bool, ocr_mode: str, render_scale: float, ocr_prompt: str):
    """
    PDF dosyasını okur. OCR modu:
      - "auto": sayfadaki çıkarılan metin boşsa (veya çok azsa) OCR uygula
      - "force": tüm sayfalara OCR uygula (çıkarılan metin olsa dahi)
    """
    data = file.getvalue() if hasattr(file, "getvalue") else file.read()
    text_parts = []
    try:
        reader = pypdf.PdfReader(BytesIO(data))
        n_pages = len(reader.pages)
    except Exception as e:
        st.error(f"PDF okunamadı: {e}")
        return "", None, None

    pdf_doc = pdfium.PdfDocument(BytesIO(data)) if ocr_enabled else None

    for i in range(n_pages):
        page = reader.pages[i]
        extracted = (page.extract_text() or "").strip()
        need_ocr = False
        if ocr_enabled:
            if ocr_mode == "force":
                need_ocr = True
            elif ocr_mode == "auto":
                # çok kısa metinleri de OCR'la
                need_ocr = len(extracted) < 30

        if need_ocr:
            pil = render_pdf_page(pdf_doc, i, scale=render_scale)
            ocr_text = qwen_ocr_image(pil, prompt=ocr_prompt)
            final_text = ocr_text.strip()
        else:
            final_text = extracted

        # Sayfa başlığı eklemek isteğe bağlı - chunking'e yardımcı olur
        text_parts.append(f"[SAYFA {i+1}]\n{final_text}\n")

    return "\n".join(text_parts), None, None

def read_docx_ocr(file, ocr_enabled: bool, ocr_prompt: str):
    """
    DOCX içindeki normal metni okur; ocr_enabled=True ise gömülü görselleri de OCR'lar.
    """
    data = file.getvalue() if hasattr(file, "getvalue") else file.read()
    try:
        d = docx.Document(BytesIO(data))
    except Exception as e:
        st.error(f"DOCX okunamadı: {e}")
        return "", None, None

    # Paragraf metinleri
    text = ""
    for p in d.paragraphs:
        text += p.text + "\n"

    if ocr_enabled:
        # Gömülü görselleri bul ve OCR yap
        # (python-docx ilişki tablosu üzerinden image rel'leri)
        rels = d.part._rels
        img_count = 0
        for rel in rels:
            rel_obj = rels[rel]
            # hedef bir resimse
            if "image" in str(rel_obj.target_ref).lower():
                try:
                    blob = rel_obj.target_part.blob
                    pil = Image.open(BytesIO(blob)).convert("RGB")
                    ocr_text = qwen_ocr_image(pil, prompt=ocr_prompt)
                    img_count += 1
                    text += f"\n[RESİM {img_count} OCR]\n{ocr_text}\n"
                except Exception:
                    # resmi açamadıysa atla
                    pass

    return text, None, None

def read_txt(file):
    file_bytes = file.read()
    enc = detect_encoding(file_bytes)
    try:
        text = file_bytes.decode(enc)
    except:
        text = file_bytes.decode('utf-8', errors='ignore')
    return text, None, None

def read_csv(file):
    try:
        df = pd.read_csv(file)
        text = df.to_string()
        columns = df.columns.tolist()
        return text, df, columns
    except Exception as e:
        st.error(f"CSV okuma hatası: {str(e)}")
        return "", None, None

def read_excel(file):
    try:
        df = pd.read_excel(file)
        text = df.to_string()
        columns = df.columns.tolist()
        return text, df, columns
    except Exception as e:
        st.error(f"Excel okuma hatası: {str(e)}")
        return "", None, None

def read_image_ocr(file, ocr_prompt: str):
    """
    Tek görsel dosyadan OCR.
    """
    try:
        data = file.getvalue() if hasattr(file, "getvalue") else file.read()
        pil = Image.open(BytesIO(data)).convert("RGB")
        text = qwen_ocr_image(pil, prompt=ocr_prompt)
        return text, None, None
    except Exception as e:
        st.error(f"Görsel okunamadı: {e}")
        return "", None, None

# =========================
# Chunking fonksiyonları
# =========================
def create_chunks_with_columns(df, columns, chunk_method, chunk_size=500, chunk_overlap=50, separator="\n\n",
                               include_columns_in_text=False):
    chunks = []
    if chunk_method == "Satır Bazlı":
        for _, row in df.iterrows():
            chunk_data = {}
            for col in columns:
                chunk_data[col] = str(row[col]) if pd.notna(row[col]) else ""
            if include_columns_in_text:
                chunk_text = ""
                for col in columns:
                    chunk_text += f"{col}: {row[col]}\n"
                chunk_data['text'] = chunk_text.strip()
            else:
                chunk_data['text'] = " ".join([str(row[col]) for col in columns if pd.notna(row[col])])
            chunks.append(chunk_data)

    elif chunk_method == "Sabit Boyut":
        current_rows = []
        current_text = ""
        for _, row in df.iterrows():
            if include_columns_in_text:
                row_text = ""
                for col in columns:
                    row_text += f"{col}: {row[col]}\n"
                row_text += "-" * 30 + "\n"
            else:
                row_text = " ".join([str(row[col]) for col in columns if pd.notna(row[col])]) + "\n"

            if len(current_text) + len(row_text) > chunk_size:
                if current_rows:
                    chunk_data = {}
                    for col in columns:
                        values = [str(r[col]) for r in current_rows if pd.notna(r[col])]
                        chunk_data[col] = values[0] if values else ""
                    chunk_data['text'] = current_text.strip()
                    chunks.append(chunk_data)
                current_rows = [row.to_dict()]
                current_text = row_text
            else:
                current_rows.append(row.to_dict())
                current_text += row_text

        if current_rows:
            chunk_data = {}
            for col in columns:
                values = [str(r[col]) for r in current_rows if pd.notna(r.get(col))]
                chunk_data[col] = values[0] if values else ""
            chunk_data['text'] = current_text.strip()
            chunks.append(chunk_data)

    else:  # Ayırıcı Bazlı
        rows_per_chunk = max(1, chunk_size // 100)
        for i in range(0, len(df), rows_per_chunk):
            chunk_rows = df.iloc[i:i + rows_per_chunk]
            chunk_data = {}
            first_row = chunk_rows.iloc[0]
            for col in columns:
                chunk_data[col] = str(first_row[col]) if pd.notna(first_row[col]) else ""
            chunk_text = ""
            for _, row in chunk_rows.iterrows():
                if include_columns_in_text:
                    for col in columns:
                        chunk_text += f"{col}: {row[col]}\n"
                    chunk_text += "-" * 30 + "\n"
                else:
                    chunk_text += " ".join([str(row[col]) for col in columns if pd.notna(row[col])]) + "\n"
            chunk_data['text'] = chunk_text.strip()
            chunks.append(chunk_data)

    return chunks

def create_chunks(text: str, method: str, chunk_size: int = 500,
                  chunk_overlap: int = 50, separator: str = "\n\n") -> List[Dict]:
    chunks = []

    if method == "Sabit Boyut":
        size = max(1, chunk_size)
        effective_overlap = min(max(0, chunk_overlap), size - 1)
        step = size - effective_overlap
        start = 0
        while start < len(text):
            end = start + size
            chunk_text = text[start:end]
            chunks.append({'text': chunk_text})
            start += step

    elif method == "Ayırıcı Bazlı":
        size = max(1, chunk_size)
        parts = text.split(separator)
        current_chunk = ""
        for part in parts:
            candidate = (current_chunk + part + separator) if current_chunk else (part + separator)
            if len(candidate) <= size:
                current_chunk = candidate
            else:
                if current_chunk:
                    chunks.append({'text': current_chunk.strip()})
                    current_chunk = part + separator
                    if len(current_chunk) > size:
                        while len(current_chunk) > size:
                            chunks.append({'text': current_chunk[:size].strip()})
                            current_chunk = current_chunk[size:]
                else:
                    piece = part + separator
                    while len(piece) > size:
                        chunks.append({'text': piece[:size].strip()})
                        piece = piece[size:]
                    current_chunk = piece
        if current_chunk:
            chunks.append({'text': current_chunk.strip()})

        if chunk_overlap > 0 and len(chunks) > 1:
            effective_overlap = min(max(0, chunk_overlap), size - 1)
            for i in range(1, len(chunks)):
                prev_text = chunks[i - 1]['text']
                prefix = prev_text[-effective_overlap:] if len(prev_text) >= effective_overlap else prev_text
                new_text = (prefix + chunks[i]['text'])
                chunks[i]['text'] = new_text[:size]

    else:  # Cümle Bazlı
        size = max(1, chunk_size)
        sentences = text.replace('!', '.').replace('?', '.').split('.')
        current_chunk = ""
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            candidate = (current_chunk + sentence + ". ") if current_chunk else (sentence + ". ")
            if len(candidate) <= size:
                current_chunk = candidate
            else:
                if current_chunk:
                    chunks.append({'text': current_chunk.strip()})
                    current_chunk = sentence + ". "
                    if len(current_chunk) > size:
                        while len(current_chunk) > size:
                            chunks.append({'text': current_chunk[:size].strip()})
                            current_chunk = current_chunk[size:]
                else:
                    piece = sentence + ". "
                    while len(piece) > size:
                        chunks.append({'text': piece[:size].strip()})
                        piece = piece[size:]
                    current_chunk = piece

        if current_chunk:
            chunks.append({'text': current_chunk.strip()})

        if chunk_overlap > 0 and len(chunks) > 1:
            effective_overlap = min(max(0, chunk_overlap), size - 1)
            for i in range(1, len(chunks)):
                prev_text = chunks[i - 1]['text']
                prefix = prev_text[-effective_overlap:] if len(prev_text) >= effective_overlap else prev_text
                new_text = (prefix + chunks[i]['text'])
                chunks[i]['text'] = new_text[:size]

    return chunks

# =========================
# Dosya işleme (OCR entegre)
# =========================
IMAGE_EXTS = ['png', 'jpg', 'jpeg', 'tif', 'tiff', 'bmp', 'webp']

def process_file(file, chunk_method, chunk_size, chunk_overlap, separator,
                 include_columns_in_text,
                 ocr_enabled: bool,
                 ocr_mode: str,
                 render_scale: float,
                 ocr_prompt: str):
    file_extension = file.name.split('.')[-1].lower()
    file_name = file.name

    if file_extension == 'pdf':
        text, df, columns = read_pdf_ocr(file, ocr_enabled, ocr_mode, render_scale, ocr_prompt)
    elif file_extension == 'docx':
        text, df, columns = read_docx_ocr(file, ocr_enabled, ocr_prompt)
    elif file_extension == 'txt':
        text, df, columns = read_txt(file)
    elif file_extension == 'csv':
        text, df, columns = read_csv(file)
    elif file_extension in ['xlsx', 'xls']:
        text, df, columns = read_excel(file)
    elif file_extension in IMAGE_EXTS:
        # tek sayfalık görsel OCR
        text, df, columns = read_image_ocr(file, ocr_prompt)
    else:
        return None, f"Desteklenmeyen dosya formatı: {file_extension}"

    if text is None:
        return None, f"Dosya okunamadı: {file_name}"

    # Chunking
    if df is not None and columns is not None:
        chunks = create_chunks_with_columns(
            df, columns, chunk_method, chunk_size, chunk_overlap,
            separator, include_columns_in_text
        )
        for chunk in chunks:
            chunk['filename'] = file_name
    else:
        chunks = create_chunks(text, chunk_method, chunk_size, chunk_overlap, separator)
        for chunk in chunks:
            chunk['filename'] = file_name

    return chunks, None

# =========================
# Görselleştirme / Dışa aktarım
# =========================
def display_chunks(all_chunks: List[Dict]):
    if not all_chunks:
        return

    st.markdown("###  Chunk İstatistikleri")

    file_stats = {}
    for chunk in all_chunks:
        filename = chunk.get('filename', 'Bilinmeyen')
        if filename not in file_stats:
            file_stats[filename] = {'count': 0, 'total_length': 0}
        file_stats[filename]['count'] += 1
        file_stats[filename]['total_length'] += len(chunk.get('text', ''))

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Toplam Chunk", len(all_chunks))
    with col2:
        st.metric("İşlenen Dosya", len(file_stats))
    with col3:
        avg_length = sum(len(c.get('text', '')) for c in all_chunks) / len(all_chunks) if all_chunks else 0
        st.metric("Ort. Uzunluk", f"{avg_length:.0f} karakter")
    with col4:
        total_chars = sum(len(c.get('text', '')) for c in all_chunks)
        st.metric("Toplam Karakter", f"{total_chars:,}")

    st.markdown("####  Dosya Detayları")
    for filename, stats in file_stats.items():
        with st.expander(f" {filename}"):
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Chunk Sayısı", stats['count'])
            with col2:
                st.metric("Toplam Karakter", f"{stats['total_length']:,}")
            with col3:
                avg = stats['total_length'] / stats['count'] if stats['count'] > 0 else 0
                st.metric("Ortalama", f"{avg:.0f} kar.")

    st.markdown("---")
    st.markdown("###  Oluşturulan Chunk'lar")

    col1, col2 = st.columns([2, 1])
    with col1:
        search_term = st.text_input(" Chunk'larda ara:", "")
    with col2:
        selected_file = st.selectbox(" Dosya filtresi:", ["Tümü"] + list(file_stats.keys()))

    col1, col2, col3 = st.columns(3)
    with col1:
        show_metadata = st.checkbox("Metadata göster", value=True)
    with col2:
        show_columns = st.checkbox("Sütunları göster", value=True)
    with col3:
        show_preview_only = st.checkbox("Sadece önizleme", value=False)

    filtered = []
    for i, chunk in enumerate(all_chunks, 1):
        text = chunk.get('text', '')
        filename = chunk.get('filename', 'Bilinmeyen')
        if search_term and search_term.lower() not in text.lower():
            continue
        if selected_file != "Tümü" and filename != selected_file:
            continue
        filtered.append((i, chunk))

    if filtered:
        st.info(f" {len(filtered)} chunk gösteriliyor")
        chunks_per_page = 10
        total_pages = (len(filtered) - 1) // chunks_per_page + 1
        page = st.slider("Sayfa", 1, total_pages, 1) if total_pages > 1 else 1
        start_idx = (page - 1) * chunks_per_page
        end_idx = min(start_idx + chunks_per_page, len(filtered))

        for idx, chunk in filtered[start_idx:end_idx]:
            with st.expander(f"Chunk {idx} - {chunk.get('filename', 'Bilinmeyen')} ({len(chunk.get('text', ''))} karakter)"):
                if show_metadata:
                    st.markdown("** Metadata:**")
                    mcols = st.columns(4)
                    with mcols[0]:
                        st.caption(f"ID: {idx}")
                    with mcols[1]:
                        st.caption(f"Dosya: {chunk.get('filename', 'N/A')}")
                    with mcols[2]:
                        st.caption(f"Uzunluk: {len(chunk.get('text', ''))}")
                    with mcols[3]:
                        st.caption(f"Kelime: {len(chunk.get('text', '').split())}")

                if show_columns:
                    coldata = {k: v for k, v in chunk.items() if k not in ['text', 'filename', 'chunk_id', 'length']}
                    if coldata:
                        st.markdown("** Sütun Verileri:**")
                        for k, v in coldata.items():
                            st.write(f"• **{k}:** {v}")

                st.markdown("** İçerik:**")
                text_to_show = chunk.get('text', '')
                if show_preview_only and len(text_to_show) > 500:
                    text_to_show = text_to_show[:500] + "..."
                st.code(text_to_show, language=None)
    else:
        st.warning("Arama kriterlerine uygun chunk bulunamadı.")

def export_chunks(chunks: List[Dict], format: str):
    if format == "JSON":
        export_data = []
        for i, chunk in enumerate(chunks, 1):
            chunk_export = {
                "chunk_id": i,
                "filename": chunk.get('filename', 'unknown')
            }
            for key, value in chunk.items():
                if key not in ['text', 'filename']:
                    chunk_export[key] = value
            chunk_export["text"] = chunk.get('text', '')
            chunk_export["length"] = len(chunk.get('text', ''))
            export_data.append(chunk_export)
        return json.dumps(export_data, ensure_ascii=False, indent=2)

    elif format == "TXT":
        output = []
        for i, chunk in enumerate(chunks, 1):
            output.append(f"=== CHUNK {i} ===")
            output.append(f"Dosya: {chunk.get('filename', 'unknown')}")
            coldata = {k: v for k, v in chunk.items() if k not in ['text', 'filename']}
            if coldata:
                output.append("Sütun Verileri:")
                for k, v in coldata.items():
                    output.append(f"  {k}: {v}")
            output.append(f"Uzunluk: {len(chunk.get('text', ''))}")
            output.append("İçerik:")
            output.append(chunk.get('text', ''))
            output.append("\n" + "=" * 50 + "\n")
        return "\n".join(output)

    elif format == "CSV":
        rows = []
        for i, chunk in enumerate(chunks, 1):
            row = {
                'chunk_id': i,
                'filename': chunk.get('filename', 'unknown'),
                'text': chunk.get('text', ''),
                'length': len(chunk.get('text', ''))
            }
            for key, value in chunk.items():
                if key not in ['text', 'filename']:
                    row[f'column_{key}'] = value
            rows.append(row)
        df = pd.DataFrame(rows)
        return df.to_csv(index=False)

# =========================
# UI - Sidebar
# =========================
with st.sidebar:
    st.header(" Chunking & OCR Ayarları")

    uploaded_files = st.file_uploader(
        "Dosya(lar) Seçin",
        type=['txt', 'pdf', 'docx', 'csv', 'xlsx'] + IMAGE_EXTS,
        accept_multiple_files=True,
        help="TXT, PDF, DOCX, CSV, XLSX ve görseller (PNG/JPG/TIFF/BMP/WEBP) desteklenir."
    )

    if uploaded_files:
        st.info(f" {len(uploaded_files)} dosya yüklendi")
        for file in uploaded_files:
            st.caption(f"• {file.name}")

    st.markdown("---")

    # OCR
    ocr_enabled = st.checkbox(" OCR (Qwen2.5-VL) kullan", value=True,
                              help="PDF/DOCX sayfa/görsellerini görüntü tabanlı modelle oku")
    ocr_mode_label = st.selectbox(
        "OCR modu",
        ["Otomatik (boş/az metinli sayfalarda)", "Zorla (tüm sayfaları OCR)"],
        index=0
    )
    ocr_mode = "auto" if "Otomatik" in ocr_mode_label else "force"

    render_scale = st.slider("PDF OCR render ölçeği", min_value=1.0, max_value=3.0, value=2.0, step=0.25,
                             help="Daha yüksek ölçek = daha net OCR (daha yavaş)")

    ocr_prompt = st.text_area(
        "OCR prompt (opsiyonel)",
        value=("Aşağıdaki görseldeki TÜM metni, doğal okuma sırası ve satır sonlarını "
               "olabildiğince koruyarak çıkar. Sadece ham metni döndür; ek açıklama, "
               "etiket veya format ekleme."),
        help="Qwen2.5-VL'ye verilecek istem. Boş bırakılırsa varsayılan kullanılır."
    )

    st.markdown("---")

    include_columns_in_text = st.checkbox(
        " Sütun başlıklarını text alanına ekle (CSV/Excel)",
        value=False,
        help="CSV/Excel dosyalarında sütun adlarını text alanına ekler. Not: Sütunlar her zaman JSON çıktısına eklenir."
    )

    chunk_method = st.selectbox(
        "Chunking Yöntemi",
        ["Sabit Boyut", "Ayırıcı Bazlı", "Cümle Bazlı", "Satır Bazlı"],
        help="Metni nasıl böleceğinizi seçin"
    )

    if chunk_method == "Sabit Boyut":
        chunk_size = st.slider("Chunk Boyutu (karakter)", 100, 2000, 500, 50)
        chunk_overlap = st.slider("Chunk Örtüşmesi (karakter)", 0, 200, 50, 10)
        separator = "\n\n"
    elif chunk_method == "Ayırıcı Bazlı":
        separator = st.text_input("Ayırıcı Karakter", value="\n\n", help="Metni bölerken kullanılacak ayırıcı")
        chunk_size = st.slider("Maksimum Chunk Boyutu", 100, 2000, 500, 50)
        chunk_overlap = st.slider("Chunk Örtüşmesi (karakter)", 0, 200, 50, 10)
    elif chunk_method == "Satır Bazlı":
        st.info("Her satır ayrı bir chunk olacak")
        chunk_size = 500
        chunk_overlap = 0
        separator = "\n"
    else:
        chunk_size = st.slider("Maksimum Chunk Boyutu", 100, 2000, 500, 50)
        chunk_overlap = st.slider("Chunk Örtüşmesi (karakter)", 0, 200, 50, 10)
        separator = "."

    st.markdown("---")
    col1, col2 = st.columns(2)
    with col1:
        process_button = st.button(" Chunk'lara Böl", type="primary", use_container_width=True)
    with col2:
        clear_button = st.button(" Temizle", type="secondary", use_container_width=True)

    if clear_button:
        st.session_state.all_chunks = []
        st.session_state.processed_files = []
        st.rerun()

    if process_button:
        if uploaded_files:
            progress_bar = st.progress(0)
            status_text = st.empty()

            all_chunks = []
            processed_files = []
            errors = []

            # Modelin önceden yüklenmesini isteğe bağlı tetikleyin (ilk OCR'da zaten cache edilecek)
            if ocr_enabled:
                try:
                    status_text.text("Qwen2.5-VL modeli yükleniyor...")
                    _ = load_qwen_model()
                except Exception as e:
                    errors.append(f"OCR modeli yüklenemedi: {e}")

            for i, file in enumerate(uploaded_files):
                status_text.text(f"İşleniyor: {file.name}")
                progress_bar.progress((i + 1) / len(uploaded_files))

                try:
                    chunks, error = process_file(
                        file, chunk_method, chunk_size, chunk_overlap,
                        separator, include_columns_in_text,
                        ocr_enabled, ocr_mode, render_scale, ocr_prompt
                    )
                    if error:
                        errors.append(error)
                    elif chunks:
                        all_chunks.extend(chunks)
                        processed_files.append(file.name)
                except Exception as e:
                    errors.append(f"{file.name}: {str(e)}")

            progress_bar.empty()
            status_text.empty()

            st.session_state.all_chunks = all_chunks
            st.session_state.processed_files = processed_files

            if all_chunks:
                st.success(f" {len(processed_files)} dosya işlendi, {len(all_chunks)} chunk oluşturuldu!")
            if errors:
                for error in errors:
                    st.error(f" {error}")
        else:
            st.warning("Lütfen en az bir dosya yükleyin!")

    if st.session_state.all_chunks:
        st.markdown("---")
        st.header(" Dışa Aktar")
        export_format = st.selectbox("Format Seçin", ["JSON", "TXT", "CSV"], help="JSON formatı tüm sütun bilgilerini içerir")
        if st.button(" İndir", use_container_width=True):
            export_data = export_chunks(st.session_state.all_chunks, export_format)
            if export_format == "JSON":
                file_name = "chunks.json"; mime = "application/json"
            elif export_format == "TXT":
                file_name = "chunks.txt"; mime = "text/plain"
            else:
                file_name = "chunks.csv"; mime = "text/csv"

            st.download_button(
                label=f" {export_format} olarak indir",
                data=export_data,
                file_name=file_name,
                mime=mime,
                use_container_width=True
            )

# =========================
# Ana içerik
# =========================
if st.session_state.all_chunks:
    display_chunks(st.session_state.all_chunks)
else:
    st.markdown("""
    ### Hoş Geldiniz!

    Bu uygulama ile:
    - **PDF/DOCX/görselleri OCR** ile okuyabilir
    - **Farklı yöntemlerle** chunk'lara bölebilir
    - **CSV/Excel sütunlarını** otomatik JSON'a ekleyebilir
    - **Sonuçları** JSON, TXT veya CSV olarak indirebilirsiniz

    **Başlamak için:**
    1. Sol panelden dosyalarınızı yükleyin
    2. OCR ve chunking ayarlarını belirleyin
    3. "Chunk'lara Böl" butonuna tıklayın
    """)
    with st.expander(" Örnek (CSV)"):
        sample_df = pd.DataFrame({
            'MüşteriID': ['C001', 'C002', 'C003'],
            'Ad': ['Ahmet', 'Ayşe', 'Mehmet'],
            'Soyad': ['Kara', 'Demir', 'Yılmaz'],
            'Şehir': ['İstanbul', 'Ankara', 'İzmir']
        })
        st.dataframe(sample_df)
        st.caption("CSV/Excel için sütunlar JSON çıktısına otomatik eklenir.")

st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #666;'>
    <small> Gelişmiş Doküman Chunking (OCR) v4.0 - Qwen/Qwen2.5-VL-7B-Instruct entegrasyonu</small>
</div>
""", unsafe_allow_html=True)

